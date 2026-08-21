"""Extracción de texto de la documentación de una licitación.

Cubre el paso 3-4 del flujo ("Descargar toda la documentación" / "Leer
íntegramente el pliego"): PDF, Word, Excel e imágenes (OCR best-effort).

Las dependencias no esenciales (python-docx, openpyxl, pytesseract) son
opcionales: si no están instaladas, el parser degrada avisando qué tipo de
archivo no pudo leer en vez de fallar todo el pipeline. Esto es intencional
— un pliego con un anexo en Excel no debería tirar abajo el análisis del
PDF principal.
"""
from __future__ import annotations

import multiprocessing
import os
import re
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, Optional

import requests

HEADERS = {"User-Agent": "Mozilla/5.0 (compatible; MetropolitanaLicitaciones/1.0)"}

EXTENSIONES_SOPORTADAS = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".png", ".jpg", ".jpeg", ".tif", ".tiff"}

# Timeout duro (en segundos) para la extracción de texto de UN documento.
#
# Evidencia real (corrida #206 de monitor.yml, 2026-08-20): un pliego en PDF
# con una fuente de codificación no estándar ("/SymbolSetEncoding") hizo que
# pypdf.extract_text() se volviera catastróficamente lento — el log repitió
# la advertencia "Advanced encoding /SymbolSetEncoding not implemented yet"
# sin parar durante 5h59m, hasta que GitHub mató el job entero por el límite
# de 6 horas. Se perdió TODA la corrida (todos los llamados ya procesados,
# no solo ese documento) porque no había ningún límite de tiempo alrededor
# de la extracción.
#
# Configurable vía env var para poder bajarlo en tests.
TIMEOUT_EXTRACCION_SEGUNDOS = int(os.environ.get("TIMEOUT_EXTRACCION_DOCUMENTO_SEGUNDOS", "120"))


@dataclass
class DocumentoExtraido:
    nombre: str
    url: str
    tipo: str
    texto: str = ""
    paginas_leidas: int = 0
    error: Optional[str] = None

    @property
    def ok(self) -> bool:
        return self.error is None and bool(self.texto.strip())


@dataclass
class PliegoExtraido:
    documentos: list[DocumentoExtraido] = field(default_factory=list)

    @property
    def texto_completo(self) -> str:
        return "\n\n".join(d.texto for d in self.documentos if d.texto)

    @property
    def documentos_con_error(self) -> list[DocumentoExtraido]:
        return [d for d in self.documentos if d.error]


# ─── Extractores por tipo de archivo ──────────────────────────────────────

def _extraer_pdf(path: Path, max_paginas: int = 40) -> tuple[str, int]:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("Falta la librería 'pypdf' (pip install pypdf)") from e

    reader = PdfReader(str(path))
    texto = []
    n = min(len(reader.pages), max_paginas)
    for i in range(n):
        texto.append(reader.pages[i].extract_text() or "")
    return "\n".join(texto), n


def _extraer_docx(path: Path) -> tuple[str, int]:
    try:
        import docx
    except ImportError as e:
        raise RuntimeError("Falta la librería 'python-docx' (pip install python-docx)") from e

    doc = docx.Document(str(path))
    partes = [p.text for p in doc.paragraphs]
    for tabla in doc.tables:
        for fila in tabla.rows:
            partes.append(" | ".join(c.text for c in fila.cells))
    return "\n".join(partes), 1


def _extraer_xlsx(path: Path) -> tuple[str, int]:
    try:
        import openpyxl
    except ImportError as e:
        raise RuntimeError("Falta la librería 'openpyxl' (pip install openpyxl)") from e

    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    partes = []
    hojas = 0
    for hoja in wb.worksheets:
        hojas += 1
        partes.append(f"### Hoja: {hoja.title}")
        for fila in hoja.iter_rows(values_only=True):
            valores = [str(v) for v in fila if v is not None]
            if valores:
                partes.append(" | ".join(valores))
    return "\n".join(partes), hojas


def _extraer_imagen(path: Path) -> tuple[str, int]:
    try:
        import pytesseract
        from PIL import Image
    except ImportError as e:
        raise RuntimeError(
            "Falta 'pytesseract'/'Pillow' o el binario de Tesseract OCR no está "
            "instalado (pip install pytesseract pillow, y el paquete de sistema tesseract-ocr)"
        ) from e

    texto = pytesseract.image_to_string(Image.open(path), lang="spa+eng")
    return texto, 1


_EXTRACTORES = {
    ".pdf": _extraer_pdf,
    ".docx": _extraer_docx,
    ".xlsx": _extraer_xlsx,
    ".png": _extraer_imagen,
    ".jpg": _extraer_imagen,
    ".jpeg": _extraer_imagen,
    ".tif": _extraer_imagen,
    ".tiff": _extraer_imagen,
}


def extraer_archivo_local(path: Path, nombre: Optional[str] = None, url: str = "") -> DocumentoExtraido:
    """Extrae texto de un archivo ya descargado en disco."""
    nombre = nombre or path.name
    ext = path.suffix.lower()
    extractor = _EXTRACTORES.get(ext)

    if extractor is None:
        if ext in (".doc", ".xls"):
            return DocumentoExtraido(
                nombre=nombre, url=url, tipo=ext,
                error=(
                    f"Formato legado '{ext}' no soportado directamente. "
                    "Convertir a .docx/.xlsx o exportar a PDF antes de analizar."
                ),
            )
        return DocumentoExtraido(nombre=nombre, url=url, tipo=ext, error=f"Extensión no soportada: {ext}")

    try:
        texto, paginas = extractor(path)
        return DocumentoExtraido(nombre=nombre, url=url, tipo=ext, texto=texto, paginas_leidas=paginas)
    except Exception as e:  # noqa: BLE001 - queremos degradar, no abortar el pipeline
        return DocumentoExtraido(nombre=nombre, url=url, tipo=ext, error=str(e))


def _correr_target_y_encolar(func, args: tuple, queue: "multiprocessing.Queue") -> None:
    """Target genérico del subproceso: corre func(*args) y pone el resultado
    en la queue. Función de módulo (no closure/lambda) para que
    multiprocessing pueda importarla en el hijo bajo el contexto "spawn".
    """
    queue.put(func(*args))


def ejecutar_con_timeout_duro(func: Callable, args: tuple, timeout_segundos: int, timeout_exitcode_a: Callable):
    """Corre func(*args) en un subproceso aparte y lo mata si no termina
    dentro de timeout_segundos.

    Usa multiprocessing (no concurrent.futures.ThreadPoolExecutor) a
    propósito: un cuelgue CPU-bound dentro de una librería en C/Python puro
    (ej. pypdf.extract_text() con una fuente de codificación rara — ver
    TIMEOUT_EXTRACCION_SEGUNDOS más arriba) no coopera con timeouts a nivel
    de threading, porque el GIL nunca se libera. Hace falta poder matar el
    proceso de verdad.

    func y sus args deben ser picklables (se pasan al subproceso vía
    multiprocessing bajo contexto "spawn"). Devuelve (resultado, se_agoto):
    si se_agoto es True, resultado es lo que devuelve timeout_exitcode_a()
    (para que el caller arme su propio objeto de error con sus propios
    campos, en vez de que esta función genérica conozca DocumentoExtraido).
    """
    ctx = multiprocessing.get_context("spawn")
    queue: "multiprocessing.Queue" = ctx.Queue()
    proceso = ctx.Process(target=_correr_target_y_encolar, args=(func, args, queue), daemon=True)
    proceso.start()
    proceso.join(timeout_segundos)

    if proceso.is_alive():
        proceso.kill()  # SIGKILL — terminate() (SIGTERM) podría no alcanzar si está trabado en código C
        proceso.join(5)
        return timeout_exitcode_a(None), True

    try:
        # get(timeout=...) en vez de get_nowait(): el feeder thread de la
        # Queue puede tardar un instante en terminar de escribir al pipe
        # después de que el proceso hijo ya se reportó como no-vivo — con
        # get_nowait() eso es una carrera real (Empty espurio).
        return queue.get(timeout=5), False
    except Exception:
        return timeout_exitcode_a(proceso.exitcode), False


def extraer_archivo_local_con_timeout(
    path: Path,
    nombre: Optional[str] = None,
    url: str = "",
    timeout_segundos: int = TIMEOUT_EXTRACCION_SEGUNDOS,
) -> DocumentoExtraido:
    """Igual que extraer_archivo_local(), pero corriendo la extracción en un
    subproceso aparte con un timeout duro.

    extraer_archivo_local() delega en librerías de terceros (pypdf,
    pytesseract, etc.) cuyo tiempo de ejecución no controlamos — un solo
    documento con una codificación/fuente problemática puede colgarse
    indefinidamente (ver comentario de TIMEOUT_EXTRACCION_SEGUNDOS más
    arriba, con la evidencia real de la corrida #206).
    """
    nombre_final = nombre or path.name
    tipo = path.suffix.lower()

    def _error_timeout(exitcode):
        if exitcode is None:
            return DocumentoExtraido(
                nombre=nombre_final, url=url, tipo=tipo,
                error=(
                    f"Extracción de texto excedió el timeout de {timeout_segundos}s "
                    "(probablemente una fuente/codificación problemática en el documento) — "
                    "se omite este documento puntual, el resto del pliego se sigue analizando."
                ),
            )
        return DocumentoExtraido(
            nombre=nombre_final, url=url, tipo=tipo,
            error=f"El subproceso de extracción terminó sin resultado (exit code {exitcode}).",
        )

    resultado, _se_agoto = ejecutar_con_timeout_duro(
        extraer_archivo_local, (path, nombre_final, url), timeout_segundos, _error_timeout,
    )
    return resultado


def descargar_y_extraer(url: str, timeout: int = 30) -> DocumentoExtraido:
    """Descarga un documento por URL y extrae su texto."""
    nombre = url.rstrip("/").split("/")[-1] or url
    ext = Path(nombre).suffix.lower()
    if not ext:
        ext = ".pdf"  # los portales de compras suelen servir PDFs sin extensión en la URL

    try:
        r = requests.get(url, headers=HEADERS, timeout=timeout)
        r.raise_for_status()
    except Exception as e:  # noqa: BLE001
        return DocumentoExtraido(nombre=nombre, url=url, tipo=ext, error=f"No se pudo descargar: {e}")

    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as f:
        f.write(r.content)
        tmp_path = Path(f.name)

    try:
        return extraer_archivo_local_con_timeout(tmp_path, nombre=nombre, url=url)
    finally:
        tmp_path.unlink(missing_ok=True)


def encontrar_links_documentos(html: str, base_url: str = "") -> list[str]:
    """Busca links a documentos descargables (pdf/doc/xls/imagen) en una página HTML."""
    patron = r'href=["\']([^"\']+\.(?:pdf|docx?|xlsx?|png|jpe?g|tiff?))["\']'
    links = re.findall(patron, html, re.IGNORECASE)
    resultado = []
    for link in links:
        if link.startswith("http"):
            resultado.append(link)
        elif base_url:
            resultado.append(base_url.rstrip("/") + "/" + link.lstrip("/"))
    return resultado


def extraer_pliego(url_licitacion: str, max_documentos: int = 10) -> PliegoExtraido:
    """Punto de entrada principal: dada la URL de una licitación en
    comprasestatales.gub.uy, descarga y extrae todos los documentos
    adjuntos que pueda encontrar.
    """
    pliego = PliegoExtraido()
    try:
        r = requests.get(url_licitacion, headers=HEADERS, timeout=20)
        r.raise_for_status()
    except Exception as e:  # noqa: BLE001
        pliego.documentos.append(
            DocumentoExtraido(nombre=url_licitacion, url=url_licitacion, tipo="pagina", error=str(e))
        )
        return pliego

    links = encontrar_links_documentos(r.text, base_url="https://www.comprasestatales.gub.uy")
    for link in links[:max_documentos]:
        pliego.documentos.append(descargar_y_extraer(link))

    return pliego
